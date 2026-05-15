"""
Abstraction de rendu rapport VAE : PDF (ReportLab) ou Word (.docx).
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Protocol, Sequence

from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import PageBreak, Paragraph, Spacer

from vae_report_style import (
    CONTENT_WIDTH,
    GREEN_LT,
    GRAY_L,
    NAVY,
    P_INS,
    P_JUST,
    P_SMALL,
    P_TOC,
    WHITE,
    alerte,
    build_doc,
    kpi_block,
    pdf_escape,
    table_standard,
    H1,
)


def strip_markup(text: str) -> str:
    t = re.sub(r"</?b>", "", str(text))
    t = t.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    return t


def _add_rich_paragraph(doc, text: str, *, style_name: str | None = None) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph(style=style_name) if style_name else doc.add_paragraph()
    parts = re.split(r"(<b>|</b>)", str(text))
    bold = False
    for part in parts:
        if part == "<b>":
            bold = True
            continue
        if part == "</b>":
            bold = False
            continue
        if not part:
            continue
        run = p.add_run(strip_markup(part))
        run.bold = bold
    if not p.runs:
        p.add_run(strip_markup(text))
    if style_name is None:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


class ReportSink(Protocol):
    width: float

    def h1(self, text: str) -> None: ...
    def paragraph(self, text: str, *, style: str = "body") -> None: ...
    def table(
        self,
        headers: Sequence[str],
        rows: Sequence[Sequence[Any]],
        *,
        col_widths: Sequence[float] | None = None,
        row_fills: Sequence[colors.Color | None] | None = None,
    ) -> None: ...
    def kpi_block(self, items: Sequence[tuple[str, str, Any | None]]) -> None: ...
    def alert(self, message: str, level: str = "info") -> None: ...
    def page_break(self) -> None: ...
    def spacer(self, gap: float = 5) -> None: ...
    def center_heading(self, text: str, *, subtitle: str | None = None) -> None: ...
    def part_banner(self, text: str) -> None: ...
    def finalize(
        self,
        path: Path,
        *,
        firm_code: str,
        footer: str,
        report_date: str | None = None,
    ) -> None: ...


class PdfSink:
    def __init__(self, story: list | None = None, width: float = CONTENT_WIDTH) -> None:
        self.story = story if story is not None else []
        self.width = width

    def h1(self, text: str) -> None:
        H1(self.story, text)

    def paragraph(self, text: str, *, style: str = "body") -> None:
        styles = {
            "body": P_INS,
            "justify": P_JUST,
            "ins": P_INS,
            "small": P_SMALL,
            "toc": P_TOC,
        }
        self.story.append(Paragraph(text, styles.get(style, P_INS)))

    def table(
        self,
        headers: Sequence[str],
        rows: Sequence[Sequence[Any]],
        *,
        col_widths: Sequence[float] | None = None,
        row_fills: Sequence[colors.Color | None] | None = None,
    ) -> None:
        w = self.width
        n = len(headers)
        if col_widths is None:
            col_widths = [w / n] * n
        self.story.append(
            table_standard(headers, rows, col_widths, row_fills=row_fills)
        )

    def kpi_block(self, items: Sequence[tuple[str, str, Any | None]]) -> None:
        self.story.append(kpi_block(items))

    def alert(self, message: str, level: str = "info") -> None:
        self.story.append(alerte(message, level))

    def page_break(self) -> None:
        self.story.append(PageBreak())

    def spacer(self, gap: float = 5) -> None:
        self.story.append(Spacer(1, gap))

    def center_heading(self, text: str, *, subtitle: str | None = None) -> None:
        hero = ParagraphStyle(
            "RAPPORT_HERO",
            parent=P_INS,
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=NAVY,
            alignment=1,
            spaceAfter=8,
        )
        sub = ParagraphStyle(
            "RAPPORT_SUB",
            parent=P_INS,
            fontSize=10,
            leading=14,
            alignment=1,
            spaceAfter=14,
            textColor=NAVY,
        )
        self.story.append(Spacer(1, 6))
        self.story.append(Paragraph(text, hero))
        if subtitle:
            self.story.append(Paragraph(pdf_escape(subtitle), sub))

    def part_banner(self, text: str) -> None:
        part = ParagraphStyle(
            "RAPPORT_PART",
            parent=P_INS,
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            textColor=NAVY,
            alignment=1,
            spaceAfter=6,
        )
        self.story.append(Paragraph(text, part))

    def finalize(
        self,
        path: Path,
        *,
        firm_code: str,
        footer: str,
        report_date: str | None = None,
    ) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        build_doc(path, self.story, firm_code=firm_code, left_text=footer, report_date=report_date)


_CM_TO_TWIPS = 567  # 1 cm ≈ 567 twips (dxa)


def _docx_usable_width_cm(doc) -> float:
    section = doc.sections[-1]
    emu = section.page_width - section.left_margin - section.right_margin
    return emu / 360000.0


def _docx_col_widths_cm(
    headers: Sequence[str],
    rows: Sequence[Sequence[Any]],
    total_cm: float,
) -> list[float]:
    n = len(headers)
    weights: list[float] = []
    for j in range(n):
        samples = [str(headers[j])] + [str(row[j]) for row in rows[:30] if j < len(row)]
        char_w = max(len(s) for s in samples)
        hdr = str(headers[j]).lower()
        if j == 0 or "scénario" in hdr or "scenario" in hdr or "note" in hdr or "description" in hdr:
            char_w = min(char_w, 28)
        elif any(k in hdr for k in ("s1", "s2", "s3", "s4", "s5", "s6", "p.", "marge", "pdm")):
            char_w = min(char_w, 10)
        weights.append(max(char_w, 5))
    total_w = sum(weights) or 1.0
    widths = [total_cm * (w / total_w) for w in weights]
    floor = total_cm / (n * 3.5)
    cap = total_cm * 0.32
    widths = [max(floor, min(w, cap)) for w in widths]
    scale = total_cm / sum(widths)
    return [w * scale for w in widths]


def _docx_oxml_replace_child(parent, tag_local: str, element) -> None:
    from docx.oxml.ns import qn

    tag = qn(f"w:{tag_local}")
    for child in list(parent):
        if child.tag == tag:
            parent.remove(child)
    parent.append(element)


def _docx_configure_table(table, width_cm: float) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_cm * _CM_TO_TWIPS)))
    _docx_oxml_replace_child(tbl_pr, "tblW", tbl_w)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    _docx_oxml_replace_child(tbl_pr, "tblLayout", layout)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    if not any(el.tag.endswith("tblHeader") for el in tr_pr):
        tr_pr.append(OxmlElement("w:tblHeader"))


def _docx_set_cell_width(cell, width_cm: float) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm

    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_cm * _CM_TO_TWIPS)))
    _docx_oxml_replace_child(tc_pr, "tcW", tc_w)


def _docx_write_cell(
    cell,
    text: str,
    *,
    header: bool = False,
    align_center: bool = False,
    font_pt: float = 8.0,
    fill: str | None = None,
) -> None:
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(max(font_pt + 2, 9))
    run = p.add_run(strip_markup(str(text)))
    run.font.name = "Calibri"
    run.font.size = Pt(font_pt)
    run.bold = header
    if header:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        fill = "1C2B4A"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (align_center or header) else WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    if fill:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        shading.set(qn("w:val"), "clear")
        _docx_oxml_replace_child(tc_pr, "shd", shading)
    tc_mar = OxmlElement("w:tcMar")
    for side, margin in (("top", "30"), ("left", "60"), ("bottom", "30"), ("right", "60")):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), margin)
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    _docx_oxml_replace_child(tc_pr, "tcMar", tc_mar)


def _docx_is_numeric_column(header: str, sample_values: Sequence[Any]) -> bool:
    hdr = header.lower()
    if any(
        k in hdr
        for k in (
            "ventes", "ca", "profit", "marge", "pdm", "service", "stock",
            "p.", "prix", "promo", "unit", "production", "période", "periode", "total",
        )
    ):
        return True
    for val in sample_values[:8]:
        raw = str(val).strip().replace(" ", "").replace(",", "").replace("%", "").replace("$", "")
        if raw and all(c.isdigit() or c in ".-" for c in raw):
            return True
    return False


class DocxSink:
    def __init__(self) -> None:
        from docx import Document
        from docx.shared import Cm, Pt

        self._Pt = Pt
        self._Cm = Cm
        self.doc = Document()
        self.width = CONTENT_WIDTH
        normal = self.doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10)
        section = self.doc.sections[0]
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.2)

    def h1(self, text: str) -> None:
        self.doc.add_heading(strip_markup(text), level=1)

    def paragraph(self, text: str, *, style: str = "body") -> None:
        if style == "small":
            p = self.doc.add_paragraph()
            run = p.add_run(strip_markup(text))
            run.font.size = self._Pt(8)
            run.italic = True
            return
        if style == "toc":
            p = self.doc.add_paragraph(style="List Bullet")
            parts = re.split(r"(<b>|</b>)", str(text))
            bold = False
            for part in parts:
                if part == "<b>":
                    bold = True
                    continue
                if part == "</b>":
                    bold = False
                    continue
                if not part:
                    continue
                run = p.add_run(strip_markup(part))
                run.bold = bold
            if not p.runs:
                p.add_run(strip_markup(text))
            return
        _add_rich_paragraph(self.doc, text)

    def table(
        self,
        headers: Sequence[str],
        rows: Sequence[Sequence[Any]],
        *,
        col_widths: Sequence[float] | None = None,
        row_fills: Sequence[colors.Color | None] | None = None,
        compact: bool = False,
    ) -> None:
        del col_widths
        ncols = len(headers)
        if ncols == 0:
            return
        width_cm = _docx_usable_width_cm(self.doc)
        widths_cm = _docx_col_widths_cm(headers, rows, width_cm)
        font_pt = 7.0 if compact or ncols >= 7 else 8.0

        table = self.doc.add_table(rows=1 + len(rows), cols=ncols)
        table.style = "Table Grid"
        _docx_configure_table(table, width_cm)

        for row in table.rows:
            for j, w_cm in enumerate(widths_cm):
                _docx_set_cell_width(row.cells[j], w_cm)

        for j, h in enumerate(headers):
            samples = [row[j] for row in rows if j < len(row)]
            center_hdr = _docx_is_numeric_column(str(h), samples)
            _docx_write_cell(
                table.rows[0].cells[j],
                str(h),
                header=True,
                align_center=center_hdr,
                font_pt=font_pt + 0.5,
            )

        for i, row in enumerate(rows):
            highlight = row_fills is not None and i < len(row_fills) and row_fills[i] is not None
            zebra = "F3F3F1" if i % 2 == 1 and not highlight else None
            for j in range(ncols):
                val = row[j] if j < len(row) else ""
                samples = [r[j] for r in rows if j < len(r)]
                center = _docx_is_numeric_column(str(headers[j]), samples)
                fill = "E8F5F0" if highlight else zebra
                _docx_write_cell(
                    table.rows[i + 1].cells[j],
                    str(val),
                    align_center=center,
                    font_pt=font_pt,
                    fill=fill,
                )

        self.spacer(6)

    def kpi_block(self, items: Sequence[tuple[str, str, Any | None]]) -> None:
        from docx.shared import Cm

        picked = list(items)[:4]
        while len(picked) < 4:
            picked.append(("", "", None))
        width_cm = _docx_usable_width_cm(self.doc)
        col_cm = width_cm / 4.0
        table = self.doc.add_table(rows=2, cols=4)
        table.style = "Table Grid"
        _docx_configure_table(table, width_cm)
        for j in range(4):
            label, val, _c = picked[j]
            for row in table.rows:
                _docx_set_cell_width(row.cells[j], col_cm)
            _docx_write_cell(table.rows[0].cells[j], label, header=True, align_center=True, font_pt=8)
            _docx_write_cell(table.rows[1].cells[j], val, align_center=True, font_pt=9, fill="EBF2FA")
        self.spacer(6)

    def alert(self, message: str, level: str = "info") -> None:
        prefix = {"critical": "CRITIQUE — ", "warning": "ATTENTION — ", "info": ""}.get(level, "")
        p = self.doc.add_paragraph()
        run = p.add_run(f"{prefix}{strip_markup(message)}")
        run.bold = level in ("critical", "warning")

    def page_break(self) -> None:
        self.doc.add_page_break()

    def spacer(self, gap: float = 5) -> None:
        del gap
        self.doc.add_paragraph()

    def center_heading(self, text: str, *, subtitle: str | None = None) -> None:
        t = self.doc.add_heading(strip_markup(text), level=0)
        t.alignment = 1
        if subtitle:
            p = self.doc.add_paragraph(strip_markup(subtitle))
            p.alignment = 1

    def part_banner(self, text: str) -> None:
        p = self.doc.add_paragraph()
        run = p.add_run(strip_markup(text))
        run.bold = True
        p.alignment = 1

    def finalize(
        self,
        path: Path,
        *,
        firm_code: str,
        footer: str,
        report_date: str | None = None,
    ) -> None:
        del firm_code
        foot = footer
        if report_date:
            foot = f"{footer} — Généré le {report_date}"
        self.doc.add_paragraph(foot)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))
