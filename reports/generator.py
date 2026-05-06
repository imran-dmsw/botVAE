"""
Report generator: produces Markdown text, Word (.docx) and PDF (.pdf) reports
from a ScenarioInput + SimulationResult pair.
"""
import io
import json
from datetime import datetime
from typing import Optional, Sequence

from engine.models import ScenarioInput, SimulationResult
from config.market_config import MARKET_CONFIG
from reporting.baseline_2026 import build_2026_baseline_summary
from reporting.recommendation_engine import build_recommendations


# ─── Markdown / text report ───────────────────────────────────────────────────


_STOCK_ALERT_TABLE = [
    ("Taux de couverture prévu", "< 90 %", "Rouge", "Risque élevé de ventes perdues : la production semble insuffisante. Risque important de ventes perdues et de perte de part de marché.", "Augmenter la production ou accepter volontairement une stratégie prudente."),
    ("Taux de couverture prévu", "90 % à 99 %", "Orange", "Production prudente : risque modéré de rupture si la demande se réalise.", "Ajuster légèrement à la hausse si l'objectif est de gagner des parts de marché."),
    ("Taux de couverture prévu", "100 % à 110 %", "Vert", "Production équilibrée : bon compromis entre disponibilité et stock.", "Maintenir la décision, sauf si le produit est très risqué ou coûteux à stocker."),
    ("Taux de couverture prévu", "111 % à 120 %", "Jaune", "Production sécuritaire : stock possible si la demande est plus faible que prévue.", "Vérifier le coût de stockage avant de confirmer."),
    ("Taux de couverture prévu", "> 120 %", "Rouge", "Risque de surproduction : stock final potentiellement élevé (coût de stockage important et risque d'invendus).", "Réduire la production, sauf stratégie volontaire de disponibilité élevée."),
    ("Ventes perdues estimées", "> 10 % de la demande prévue", "Orange/Rouge", "Risque important de ventes perdues et de perte de part de marché.", "Augmenter la production si la marge unitaire est attractive."),
]


def _stock_short_messages(result: SimulationResult) -> list[str]:
    msgs: list[str] = []
    cov = result.forecast_coverage_rate
    end_pct = (result.forecast_ending_stock_units / result.demand) if result.demand > 0 else 0.0
    gross_profit = max(result.revenue - result.production_cost, 0.0)
    storage_vs_gp = (result.inventory_carrying_cost / gross_profit) if gross_profit > 0 else 0.0
    if cov < 0.90:
        msgs.append("Alerte rupture : votre production ne couvre pas suffisamment la demande prévue.")
    elif cov < 1.00:
        msgs.append("Production prudente : vous limitez le stock, mais vous risquez des ventes perdues.")
    elif cov <= 1.10:
        msgs.append("Production équilibrée : votre décision couvre la demande avec une marge raisonnable.")
    if end_pct > 0.20:
        msgs.append("Alerte surstock : votre production dépasse fortement la demande prévue.")
    elif end_pct > 0.10:
        msgs.append("Attention : votre décision peut créer du surstockage et des coûts de stockage.")
    if storage_vs_gp > 0.05:
        msgs.append("Attention : le stockage risque de réduire fortement votre profit.")
    return msgs


def _pdf_safe(text: object) -> str:
    """
    Normalize text for core FPDF fonts (latin-1).
    Replaces common unicode punctuation and degrades unsupported chars safely.
    """
    s = str(text)
    s = (
        s.replace("–", "-")
        .replace("—", "-")
        .replace("’", "'")
        .replace("“", '"')
        .replace("”", '"')
        .replace("…", "...")
    )
    return s.encode("latin-1", "replace").decode("latin-1")

def generate_markdown_report(scenario: ScenarioInput, result: SimulationResult) -> str:
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    seg_label = MARKET_CONFIG["segments"][scenario.segment]["label"]
    range_label = MARKET_CONFIG["ranges"][scenario.model_range]["label"]
    mkt_max = scenario.adjusted_budget * MARKET_CONFIG["constraints"]["marketing_max_pct"]
    rd_max = scenario.adjusted_budget * MARKET_CONFIG["constraints"]["rd_max_pct"]

    status_icon = "✅" if result.is_valid else "❌"
    baseline_2026 = build_2026_baseline_summary(scenario, result)

    lines = [
        f"# Rapport de simulation VAE",
        f"**Généré le :** {now}  |  **Statut :** {status_icon} {'Valide' if result.is_valid else 'Non valide'}",
        "",
        "---",
        "",
        "## 1. Résumé exécutif",
        "",
        f"| Indicateur | Valeur |",
        f"|---|---|",
        f"| Firme | {result.firm_name} |",
        f"| Période | {result.period} |",
        f"| Scénario | {result.scenario_name} |",
        f"| Ventes | {result.sales:,} unités |",
        f"| Chiffre d'affaires | {result.revenue:,.0f} $ |",
        f"| Profit | {result.profit:,.0f} $ |",
        f"| Profit rate | {result.profit_rate*100:.2f}% ({result.profit_rate_status}) |",
        f"| Marge | {result.margin*100:.1f}% |",
        f"| Part de marché totale | {result.market_share*100:.2f}% |",
        f"| Part de marché segment ({seg_label}) | {result.market_share_segment*100:.2f}% |",
        f"| Taux de service | {result.service_rate*100:.1f}% |",
        f"| Score innovation | {result.innovation_score:.1f}/10 |",
        f"| Score durabilité | {result.sustainability_score:.1f}/10 |",
        "",
        "---",
        "",
        "## 2. Hypothèses du scénario",
        "",
        f"- **Modèle :** {scenario.model_name} ({range_label}, segment {seg_label})",
        f"- **Statut produit :** {scenario.product_status}",
        f"- **Prix :** {scenario.price:,.0f} $ (promotion : {scenario.promotion_rate*100:.1f}%)",
        f"- **Production :** {scenario.production:,} unités",
        f"- **Budget marketing :** {scenario.marketing_budget:,.0f} $ "
        f"({scenario.marketing_budget/mkt_max*100:.1f}% du plafond)",
        f"  - Digital : {scenario.marketing_channels.digital:,.0f} $",
        f"  - Réseaux sociaux : {scenario.marketing_channels.social_media:,.0f} $",
        f"  - Influenceurs : {scenario.marketing_channels.influencers:,.0f} $",
        f"  - Affichage : {scenario.marketing_channels.display:,.0f} $",
        f"  - Événements : {scenario.marketing_channels.events:,.0f} $",
        f"- **Budget R&D :** {scenario.rd_budget:,.0f} $ "
        f"({scenario.rd_budget/rd_max*100:.1f}% du plafond) | Projets : {scenario.rd_projects}",
        f"- **Investissement durabilité :** {scenario.sustainability_investment:,.0f} $",
        f"- **Budget ajusté de référence :** {scenario.adjusted_budget:,.0f} $",
        "",
        "---",
        "",
        "## 3. Résultats financiers détaillés",
        "",
        f"| Poste | Montant (CAD) | % du CA |",
        f"|---|---|---|",
        f"| Chiffre d'affaires | {result.revenue:,.0f} $ | 100.0% |",
        f"| Coûts de production | -{result.production_cost:,.0f} $ | {result.production_cost/max(result.revenue,1)*100:.1f}% |",
        f"| Coûts de distribution | -{result.distribution_cost:,.0f} $ | {result.distribution_cost/max(result.revenue,1)*100:.1f}% |",
        f"| Coûts marketing | -{result.marketing_cost:,.0f} $ | {result.marketing_cost/max(result.revenue,1)*100:.1f}% |",
        f"| Coûts R&D | -{result.rd_cost:,.0f} $ | {result.rd_cost/max(result.revenue,1)*100:.1f}% |",
        f"| Coûts d'exploitation | -{result.operating_cost:,.0f} $ | {result.operating_cost/max(result.revenue,1)*100:.1f}% |",
        f"| SAV / Garantie | -{result.aftersales_cost:,.0f} $ | {result.aftersales_cost/max(result.revenue,1)*100:.1f}% |",
        f"| Durabilité | -{result.sustainability_cost:,.0f} $ | {result.sustainability_cost/max(result.revenue,1)*100:.1f}% |",
        f"| **Total coûts** | **-{result.total_cost:,.0f} $** | **{result.total_cost/max(result.revenue,1)*100:.1f}%** |",
        f"| **Profit** | **{result.profit:,.0f} $** | **{result.margin*100:.1f}%** |",
        "",
        "---",
        "",
        "## 4. Alertes",
        "",
    ]

    if result.alerts:
        for alert in result.alerts:
            lines.append(f"- ⚠️ {alert}")
    else:
        lines.append("Aucune alerte — scénario conforme aux règles métier.")

    lines += [
        "",
        "---",
        "",
        "## 5. Controles metier",
        "",
        f"- Cohérence prix/gamme : **{result.price_range_consistency_status}**",
        f"- Promo valide : **{'oui' if scenario.promotion_rate in (0.0, -0.02, -0.03, -0.04, -0.05, -0.10) else 'non'}**",
        f"- Profit cible 5-10% : **{'oui' if 0.05 <= result.profit_rate <= 0.10 else 'non'}**",
        f"- Efficacite marketing : **{result.marketing_efficiency:.4f} ventes/$**",
        f"- Efficacite production : **{result.production_efficiency*100:.1f}%**",
        f"- Production recommandee N+1 : **{result.next_period_recommended_production:,}**",
        f"- Limite retrait produit : **{result.withdrawal_limit_status}**",
        "",
        "---",
        "",
        "## 6. Production, demande et gestion du stock",
        "",
        "| Indicateur | Valeur scénario |",
        "|---|---|",
        f"| Stock disponible prévu | {result.stock_available_units:,} unités |",
        f"| Taux de couverture prévu | {result.forecast_coverage_rate*100:.1f}% |",
        f"| Ventes perdues estimées | {result.forecast_lost_sales_units:,.1f} unités |",
        f"| Stock final prévu | {result.forecast_ending_stock_units:,} unités |",
        f"| Coût de stockage estimé | {result.inventory_carrying_cost:,.0f} $ |",
        "",
        "### Tableau principal des seuils d'alerte",
        "",
        "| Indicateur | Seuil | Niveau d'alerte | Message à afficher à l'étudiant | Décision possible |",
        "|---|---|---|---|---|",
    ]
    for ind, thr, lvl, msg, decision in _STOCK_ALERT_TABLE:
        lines.append(f"| {ind} | {thr} | {lvl} | {msg} | {decision} |")
    lines += [
        "",
        "### Messages d'alerte à intégrer dans la feuille de résultats",
        "",
    ]
    for msg in _stock_short_messages(result):
        lines.append(f"- {msg}")
    lines += [
        "",
        "---",
        "",
        "## 7. Analyse et interprétation",
        "",
    ]

    for interp in result.interpretations:
        lines.append(f"- {interp}")

    lines += [
        "",
        "---",
        "",
        "## 8. Reference 2026",
        "",
        f"- Taille du marche 2026: **{baseline_2026['market_size_2026']:,}** unites",
        f"- Baseline prix 2026: **{baseline_2026['baseline_price_2026']:,.0f} $**",
        f"- Baseline part de marche firme: **{baseline_2026['baseline_market_share_2026']*100:.2f}%**",
        f"- Baseline rentabilite: **{baseline_2026['baseline_profitability_2026']*100:.2f}%**",
        f"- Delta CA scenario vs baseline: **{baseline_2026['scenario_vs_baseline_revenue_delta']:,.0f} $**",
        "",
        "---",
        "",
        "## 9. Recommandations",
        "",
    ]
    for rec in build_recommendations(scenario, result):
        lines.append(f"- {rec}")

    lines += [
        "",
        "---",
        f"*Rapport généré automatiquement par le Bot de Simulation VAE — {now}*",
    ]

    return "\n".join(lines)


def _generate_recommendations(scenario: ScenarioInput, result: SimulationResult):
    recs = []
    cfg = MARKET_CONFIG["constraints"]
    mkt_max = scenario.adjusted_budget * cfg["marketing_max_pct"]
    rd_max = scenario.adjusted_budget * cfg["rd_max_pct"]

    if result.margin < cfg["min_profit_rate"]:
        recs.append(
            "**Améliorer la rentabilité** : augmenter le prix, réduire les coûts variables "
            "ou diminuer les dépenses marketing/R&D pour atteindre le seuil minimal de 2 %."
        )
    if result.service_rate < 0.85:
        recs.append(
            f"**Augmenter la production** : le taux de service est de {result.service_rate*100:.0f}%. "
            f"Envisager une production de {int(result.demand * 1.05):,} unités."
        )
    if scenario.marketing_budget < mkt_max * 0.3 and result.market_share_segment < 0.10:
        recs.append(
            "**Investir davantage en marketing** : la part de marché est faible et le budget marketing "
            "est bien en dessous du plafond autorisé."
        )
    if scenario.rd_budget == 0:
        recs.append(
            "**Allouer un budget R&D** : sans investissement R&D, le score innovation diminue chaque période, "
            "réduisant l'attractivité future du produit."
        )
    if result.market_share_segment > 0.30:
        recs.append(
            "**Position dominante** : surveiller la réaction des concurrents et consolider l'avantage "
            "par l'innovation et la durabilité."
        )
    if not recs:
        recs.append("Le scénario est globalement satisfaisant. Maintenir la stratégie actuelle.")

    return [f"- {r}" for r in recs]


def _strip_md_bold(text: str) -> str:
    return str(text).replace("**", "").strip()


def build_scenario_description(scenario: ScenarioInput, result: SimulationResult) -> str:
    """Paragraphe de synthèse pédagogique (PDF / narrative)."""
    seg = MARKET_CONFIG["segments"][scenario.segment]["label"]
    rng = MARKET_CONFIG["ranges"][scenario.model_range]["label"]
    sn = result.scenario_name or "Scenario"
    eff_price = scenario.price * (1.0 + scenario.promotion_rate)
    chunks = [
        f"Ce scenario « {sn} » place la firme {result.firm_name} en periode {result.period}. "
        f"Le produit {scenario.model_name} vise le segment « {seg} » en gamme « {rng} ». "
        f"Prix catalogue {scenario.price:,.0f} $, promotion {scenario.promotion_rate*100:.1f} %, "
        f"prix net {eff_price:,.0f} $, production planifiee {scenario.production:,} unites.",
        f"Resultats : {result.sales:,} unites vendues, CA {result.revenue:,.0f} $, "
        f"profit {result.profit:,.0f} $, marge {result.margin*100:.1f} %. "
        f"PDM segment {result.market_share_segment*100:.2f} %, service {result.service_rate*100:.0f} %.",
    ]
    if scenario.liquidation:
        chunks.append("Mode liquidation active : anticiper production nulle en periode suivante.")
    if scenario.new_model_launch or scenario.product_status == "pre_launch":
        chunks.append("Produit nouveau ou pre-lancement : plafonds et diffusion premiere annee a respecter.")
    if scenario.withdraw_model:
        chunks.append("Retrait produit envisage : verifier limites de simulation.")
    pr = result.profit_rate
    if pr < 0.05:
        chunks.append("Rentabilite sous la zone cible usuelle (5-10 % du CA).")
    elif pr > 0.10:
        chunks.append("Rentabilite elevee : marge confortable, opportunite de croissance en volume ou en part.")
    if result.service_rate < 0.9 and result.demand > result.sales:
        chunks.append("Ecart demande-ventes : risque de penurie ou de file d'attente cote client.")
    return " ".join(chunks)


def merge_improvement_advice(scenario: ScenarioInput, result: SimulationResult) -> list[str]:
    """Fusionne moteur de recommandations et heuristiques internes, sans doublons."""
    seen: set[str] = set()
    out: list[str] = []
    for rec in build_recommendations(scenario, result):
        t = _strip_md_bold(rec)
        k = t.lower()[:120]
        if k and k not in seen:
            seen.add(k)
            out.append(t)
    for line in _generate_recommendations(scenario, result):
        t = _strip_md_bold(line.lstrip("- "))
        k = t.lower()[:120]
        if k and k not in seen:
            seen.add(k)
            out.append(t)
    if not out:
        out.append("Poursuivre le suivi des indicateurs profit, taux de service et part de marche segment.")
    return out


# ─── JSON export ──────────────────────────────────────────────────────────────

def generate_json_report(scenario: ScenarioInput, result: SimulationResult) -> str:
    data = {
        "generated_at": datetime.now().isoformat(),
        "scenario": scenario.model_dump(),
        "result": result.model_dump(),
    }
    return json.dumps(data, ensure_ascii=False, indent=2, default=str)


# ─── Word (.docx) report ──────────────────────────────────────────────────────

def generate_word_report(scenario: ScenarioInput, result: SimulationResult) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    seg_label = MARKET_CONFIG["segments"][scenario.segment]["label"]
    range_label = MARKET_CONFIG["ranges"][scenario.model_range]["label"]

    # Title
    title = doc.add_heading("Rapport de Simulation VAE", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"Firme : {result.firm_name}  |  Période : {result.period}  |  Généré le : {now}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    status = "✅ Scénario valide" if result.is_valid else "❌ Scénario non valide"
    p = doc.add_paragraph(status)
    run = p.runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00) if result.is_valid else RGBColor(0xCC, 0x00, 0x00)

    doc.add_heading("1. Résumé exécutif", level=1)

    kpis = [
        ("Ventes", f"{result.sales:,} unités"),
        ("Chiffre d'affaires", f"{result.revenue:,.0f} $"),
        ("Profit", f"{result.profit:,.0f} $"),
        ("Marge", f"{result.margin*100:.1f}%"),
        ("Part de marché totale", f"{result.market_share*100:.2f}%"),
        (f"Part de marché {seg_label}", f"{result.market_share_segment*100:.2f}%"),
        ("Taux de service", f"{result.service_rate*100:.1f}%"),
        ("Score innovation", f"{result.innovation_score:.1f}/10"),
        ("Score durabilité", f"{result.sustainability_score:.1f}/10"),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Indicateur"
    hdr[1].text = "Valeur"
    for k, v in kpis:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v

    doc.add_heading("2. Résultats financiers", level=1)

    fin_data = [
        ("Chiffre d'affaires", result.revenue),
        ("Coûts de production", -result.production_cost),
        ("Coûts de distribution", -result.distribution_cost),
        ("Coûts marketing", -result.marketing_cost),
        ("Coûts R&D", -result.rd_cost),
        ("Coûts d'exploitation", -result.operating_cost),
        ("SAV / Garantie", -result.aftersales_cost),
        ("Durabilité", -result.sustainability_cost),
        ("Profit", result.profit),
    ]

    fin_table = doc.add_table(rows=1, cols=2)
    fin_table.style = "Table Grid"
    hdr2 = fin_table.rows[0].cells
    hdr2[0].text = "Poste"
    hdr2[1].text = "Montant (CAD)"
    for label, val in fin_data:
        row = fin_table.add_row().cells
        row[0].text = label
        row[1].text = f"{val:,.0f} $"

    doc.add_heading("3. Alertes", level=1)
    if result.alerts:
        for a in result.alerts:
            doc.add_paragraph(f"• {a}")
    else:
        doc.add_paragraph("Aucune alerte — scénario conforme.")

    doc.add_heading("4. Interprétation", level=1)
    for i in result.interpretations:
        doc.add_paragraph(f"• {i}")

    doc.add_heading("5. Recommandations", level=1)
    recs = _generate_recommendations(scenario, result)
    for r in recs:
        doc.add_paragraph(r.lstrip("- "))

    doc.add_paragraph(f"\nRapport généré automatiquement par le Bot de Simulation VAE — {now}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── PDF report ───────────────────────────────────────────────────────────────

def generate_pdf_report(scenario: ScenarioInput, result: SimulationResult) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=10)

    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    seg_label = MARKET_CONFIG["segments"][scenario.segment]["label"]

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Rapport de Simulation VAE", ln=True, align="C")
    pdf.set_font("Helvetica", size=9)
    pdf.cell(
        0,
        6,
        _pdf_safe(f"Firme : {result.firm_name}  |  Periode : {result.period}  |  {now}"),
        ln=True,
        align="C",
    )
    status_txt = "Scenario VALIDE" if result.is_valid else "Scenario NON VALIDE"
    pdf.set_text_color(0, 128, 0) if result.is_valid else pdf.set_text_color(200, 0, 0)
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(0, 8, status_txt, ln=True, align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    # Section 1
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "1. Résumé exécutif", ln=True)
    pdf.set_font("Helvetica", size=10)

    kpis = [
        ("Ventes", f"{result.sales:,} unites"),
        ("Chiffre d'affaires", f"{result.revenue:,.0f} $"),
        ("Profit", f"{result.profit:,.0f} $"),
        ("Marge", f"{result.margin*100:.1f}%"),
        ("Part marche totale", f"{result.market_share*100:.2f}%"),
        (f"Part marche {seg_label}", f"{result.market_share_segment*100:.2f}%"),
        ("Taux de service", f"{result.service_rate*100:.1f}%"),
        ("Score innovation", f"{result.innovation_score:.1f}/10"),
        ("Score durabilite", f"{result.sustainability_score:.1f}/10"),
    ]
    # usable width = 210 - 2*15 = 180mm; split 60/40
    cw1, cw2 = 110, 70
    for k, v in kpis:
        pdf.cell(cw1, 7, _pdf_safe(k), border=1)
        pdf.cell(cw2, 7, _pdf_safe(v), border=1, ln=True)
    pdf.ln(4)

    # Section 2
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "2. Resultats financiers", ln=True)
    pdf.set_font("Helvetica", size=10)

    fin_rows = [
        ("Chiffre d'affaires", f"{result.revenue:,.0f} $"),
        ("Couts production", f"-{result.production_cost:,.0f} $"),
        ("Couts distribution", f"-{result.distribution_cost:,.0f} $"),
        ("Couts marketing", f"-{result.marketing_cost:,.0f} $"),
        ("Couts R&D", f"-{result.rd_cost:,.0f} $"),
        ("Couts exploitation", f"-{result.operating_cost:,.0f} $"),
        ("SAV/Garantie", f"-{result.aftersales_cost:,.0f} $"),
        ("Durabilite", f"-{result.sustainability_cost:,.0f} $"),
        ("PROFIT", f"{result.profit:,.0f} $"),
    ]
    for k, v in fin_rows:
        pdf.cell(cw1, 7, _pdf_safe(k), border=1)
        pdf.cell(cw2, 7, _pdf_safe(v), border=1, ln=True)
    pdf.ln(4)

    page_w = pdf.w - pdf.l_margin - pdf.r_margin

    # Section 3 — Alerts
    pdf.set_x(pdf.l_margin)
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "3. Alertes", ln=True)
    pdf.set_font("Helvetica", size=10)
    if result.alerts:
        for a in result.alerts:
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(page_w, 6, _pdf_safe(f"- {a}"))
    else:
        pdf.set_x(pdf.l_margin)
        pdf.cell(0, 6, "Aucune alerte.", ln=True)
    pdf.ln(2)

    # Section 4 — Interpretations
    pdf.set_x(pdf.l_margin)
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "4. Interpretation", ln=True)
    pdf.set_font("Helvetica", size=10)
    for i in result.interpretations:
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(page_w, 6, _pdf_safe(f"- {i}"))
    pdf.ln(2)

    # Section 5 — Recommendations
    pdf.set_x(pdf.l_margin)
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "5. Recommandations", ln=True)
    pdf.set_font("Helvetica", size=10)
    recs = _generate_recommendations(scenario, result)
    for r in recs:
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(page_w, 6, _pdf_safe(r.lstrip("- ")))
    pdf.ln(2)

    # Section 6 — Stock thresholds (document pedagogique)
    pdf.set_x(pdf.l_margin)
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "6. Production, demande et gestion du stock", ln=True)
    pdf.set_font("Helvetica", size=10)
    stock_rows = [
        f"Stock disponible prevu: {result.stock_available_units:,} unites",
        f"Taux de couverture prevu: {result.forecast_coverage_rate*100:.1f}%",
        f"Ventes perdues estimees: {result.forecast_lost_sales_units:,.1f} unites",
        f"Stock final prevu: {result.forecast_ending_stock_units:,} unites",
        f"Cout de stockage estime: {result.inventory_carrying_cost:,.0f} $",
    ]
    for row in stock_rows:
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(page_w, 6, _pdf_safe(f"- {row}"))
    pdf.ln(1)
    for ind, thr, lvl, msg, decision in _STOCK_ALERT_TABLE:
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(
            page_w,
            5.5,
            _pdf_safe(f"[{lvl}] {ind} {thr} -> {msg} Decision: {decision}"),
        )
    for msg in _stock_short_messages(result):
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(page_w, 5.5, _pdf_safe(f"- {msg}"))

    pdf.set_font("Helvetica", "I", 8)
    pdf.ln(4)
    pdf.cell(
        0,
        6,
        _pdf_safe(f"Rapport genere automatiquement par le Bot de Simulation VAE - {now}"),
        ln=True,
        align="C",
    )

    return bytes(pdf.output())


def generate_multi_pdf_report(
    scenarios: Sequence[ScenarioInput],
    results: Sequence[SimulationResult],
) -> bytes:
    """Build a single PDF that consolidates multiple simulated scenarios."""
    from fpdf import FPDF

    if len(scenarios) != len(results):
        raise ValueError("Le nombre de scenarios et de resultats doit etre identique.")
    if not scenarios:
        raise ValueError("Aucun scenario a exporter.")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=10)
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    page_w = pdf.w - pdf.l_margin - pdf.r_margin

    def _safe_multiline(text: object, h: float = 6) -> None:
        """Write text with explicit usable width to avoid horizontal-space errors."""
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(page_w, h, _pdf_safe(text))

    # Cover
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Rapport Multi-Scenarios VAE", ln=True, align="C")
    pdf.set_font("Helvetica", size=9)
    pdf.cell(0, 6, _pdf_safe(f"Genere le : {now}"), ln=True, align="C")
    pdf.cell(0, 6, _pdf_safe(f"Nombre de scenarios : {len(results)}"), ln=True, align="C")
    pdf.ln(6)

    # Consolidated comparison block (text lines instead of fixed table to avoid
    # horizontal-space rendering issues in FPDF with variable content widths).
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "1. Tableau comparatif", ln=True)
    pdf.set_font("Helvetica", size=10)
    _safe_multiline("Colonnes: Scenario | Ventes | Profit ($) | Marge | PDM | Service")
    pdf.ln(1)
    for r in results:
        line = (
            f"- {(r.scenario_name or 'Scenario')[:40]} | "
            f"{r.sales:,} | {r.profit:,.0f} $ | {r.margin*100:.1f}% | "
            f"{r.market_share*100:.2f}% | {r.service_rate*100:.1f}%"
        )
        _safe_multiline(line)
    pdf.ln(4)

    # Best scenarios
    best_profit = max(results, key=lambda x: x.profit)
    best_margin = max(results, key=lambda x: x.margin)
    best_pdm = max(results, key=lambda x: x.market_share)
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 7, "2. Top performances", ln=True)
    pdf.set_font("Helvetica", size=10)
    _safe_multiline(f"- Meilleur profit : {best_profit.scenario_name} ({best_profit.profit:,.0f} $)")
    _safe_multiline(f"- Meilleure marge : {best_margin.scenario_name} ({best_margin.margin*100:.1f}%)")
    _safe_multiline(f"- Meilleure PDM : {best_pdm.scenario_name} ({best_pdm.market_share*100:.2f}%)")

    # Detailed section per scenario
    for i, (sc, res) in enumerate(zip(scenarios, results), start=1):
        pdf.add_page()
        page_w = pdf.w - pdf.l_margin - pdf.r_margin
        seg_label = MARKET_CONFIG["segments"][sc.segment]["label"]

        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 8, _pdf_safe(f"Scenario {i} - {res.scenario_name}"), ln=True)
        pdf.set_font("Helvetica", size=10)
        pdf.cell(
            0,
            6,
            _pdf_safe(f"Firme : {res.firm_name}  |  Periode : {res.period}  |  Segment : {seg_label}"),
            ln=True,
        )
        pdf.ln(2)

        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7, "Hypotheses", ln=True)
        pdf.set_font("Helvetica", size=10)
        _safe_multiline(
            f"Prix: {sc.price:,.0f} $ | Promotion: {sc.promotion_rate*100:.1f}% | Production: {sc.production:,} unites"
        )
        _safe_multiline(
            f"Marketing: {sc.marketing_budget:,.0f} $ | R&D: {sc.rd_budget:,.0f} $ | Durabilite: {sc.sustainability_investment:,.0f} $"
        )
        pdf.ln(1)

        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7, "KPIs", ln=True)
        pdf.set_font("Helvetica", size=10)
        _safe_multiline(
            f"Ventes: {res.sales:,} | CA: {res.revenue:,.0f} $ | Profit: {res.profit:,.0f} $ | "
            f"Marge: {res.margin*100:.1f}% | PDM: {res.market_share*100:.2f}% | Service: {res.service_rate*100:.1f}%"
        )
        pdf.ln(1)

        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7, "Production, demande et stock", ln=True)
        pdf.set_font("Helvetica", size=10)
        _safe_multiline(
            f"Stock dispo: {res.stock_available_units:,} | Couverture: {res.forecast_coverage_rate*100:.1f}% | "
            f"Ventes perdues: {res.forecast_lost_sales_units:,.1f} | Stock final: {res.forecast_ending_stock_units:,} | "
            f"Cout stockage: {res.inventory_carrying_cost:,.0f} $"
        )
        for msg in _stock_short_messages(res):
            _safe_multiline(f"- {msg}", 5.5)
        pdf.ln(1)

        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7, "Description", ln=True)
        pdf.set_font("Helvetica", size=10)
        _safe_multiline(build_scenario_description(sc, res), 5.5)
        pdf.ln(1)

        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7, "Conseils d'amelioration", ln=True)
        pdf.set_font("Helvetica", size=10)
        for tip in merge_improvement_advice(sc, res):
            _safe_multiline(f"- {tip}", 5.5)
        pdf.ln(1)

        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7, "Alertes", ln=True)
        pdf.set_font("Helvetica", size=10)
        if res.alerts:
            for alert in res.alerts:
                _safe_multiline(f"- {alert}")
        else:
            pdf.cell(0, 6, "Aucune alerte.", ln=True)

    pdf.set_font("Helvetica", "I", 8)
    pdf.ln(4)
    pdf.cell(
        0,
        6,
        _pdf_safe(f"Rapport genere automatiquement par le Bot de Simulation VAE - {now}"),
        ln=True,
        align="C",
    )
    return bytes(pdf.output())
